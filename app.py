import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

# --- 1. CONFIGURACIÓN ---
st.set_page_config(page_title="ExamenMatic AI", page_icon="🧠", layout="centered")

# --- 2. GESTIÓN DE LA API KEY (AUTOMÁTICA) ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
except FileNotFoundError:
    st.error("⚠️ No se encontró la API Key. Configura los 'Secrets' en Streamlit Cloud.")
    st.stop()

# --- 3. FUNCIONES DE LECTURA ---
def extract_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages: text += page.extract_text() or ""
    return text

def extract_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs: text += para.text + "\n"
    return text

def extract_from_txt(file): return file.getvalue().decode("utf-8")

# --- 4. FUNCIONES DE EXPORTACIÓN (MODIFICADAS) ---

def create_word_file(exam_text, sol_text):
    doc = docx.Document()
    
    # PARTE 1: EL EXAMEN
    doc.add_heading('EXAMEN', 0)
    for p in exam_text.split('\n'): 
        if p.strip(): doc.add_paragraph(p)
    
    # SALTO DE PÁGINA OBLIGATORIO
    doc.add_page_break()
    
    # PARTE 2: LAS RESPUESTAS
    doc.add_heading('SOLUCIONARIO', 0)
    for p in sol_text.split('\n'): 
        if p.strip(): doc.add_paragraph(p)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_file(exam_text, sol_text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 50
    
    # Función auxiliar para escribir texto en el PDF
    def write_text_block(text, title):
        c.setFont("Helvetica-Bold", 16)
        c.drawString(margin, height - 50, title)
        c.setFont("Helvetica", 12)
        y = height - 80
        
        for line in text.split('\n'):
            wrapped_lines = simpleSplit(line, "Helvetica", 12, width - 2*margin)
            for wrapped_line in wrapped_lines:
                if y < 50: # Si se acaba la hoja
                    c.showPage() # Nueva hoja
                    c.setFont("Helvetica", 12)
                    y = height - 50
                c.drawString(margin, y, wrapped_line)
                y -= 15
            y -= 5
        return y

    # Escribir Examen
    write_text_block(exam_text, "EXAMEN")
    
    # FORZAR NUEVA PÁGINA PARA EL SOLUCIONARIO
    c.showPage()
    
    # Escribir Solucionario
    write_text_block(sol_text, "SOLUCIONARIO")
    
    c.save()
    buffer.seek(0)
    return buffer

# --- 5. INTELIGENCIA ARTIFICIAL ---
def generate_quiz(text, num, level):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""
        Actúa como profesor experto.
        1. Crea {num} preguntas de nivel {level} sobre el texto. NO marques la respuesta correcta aquí.
        2. Escribe una línea exacta: "---SOLUCIONARIO---"
        3. Escribe las respuestas correctas con justificación.
        Texto: {text[:12000]}
        """
        return model.generate_content(prompt).text
    except Exception as e: return f"Error: {e}"

# --- 6. INTERFAZ ---
if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0
def reset_app(): st.session_state.uploader_key += 1

with st.sidebar:
    try: st.image("logo.png", width=200)
    except: st.header("🤖 ExamenMatic")
    
    st.success("✅ Sistema Conectado")
    
    difficulty = st.selectbox("Nivel:", ["Fácil", "Intermedio", "Difícil"])
    num_questions = st.slider("Preguntas:", 5, 20, 10)
    st.markdown("---")
    if st.button("🔄 Reiniciar App", on_click=reset_app): st.write("...")

st.title("🎓 Generador de Exámenes")
st.markdown("Sube tu archivo. Obtendrás un único documento con el examen y las respuestas al final.")

uploaded_file = st.file_uploader("Archivo (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key=str(st.session_state.uploader_key))

if uploaded_file:
    if st.button("🚀 Generar Examen"):
        with st.spinner("Creando evaluación..."):
            ftype = uploaded_file.name.split('.')[-1].lower()
            raw = ""
            try:
                if ftype == 'pdf': raw = extract_from_pdf(uploaded_file)
                elif ftype == 'docx': raw = extract_from_docx(uploaded_file)
                elif ftype == 'txt': raw = extract_from_txt(uploaded_file)
            except: st.error("Error leyendo archivo"); st.stop()

            if len(raw) > 50:
                res = generate_quiz(raw, num_questions, difficulty)
                
                if "---SOLUCIONARIO---" in res:
                    parts = res.split("---SOLUCIONARIO---")
                    exam, sol = parts[0], parts[1]
                else: exam, sol = res, "Error separando respuestas."

                st.success("¡Documento Generado!")
                
                # Vista previa en pestañas (para que lo veas antes de descargar)
                t1, t2 = st.tabs(["Vista Examen", "Vista Solucionario"])
                with t1: st.write(exam)
                with t2: st.write(sol)
                
                st.markdown("### 📥 Descargar Documento Completo")
                col1, col2 = st.columns(2)
                
                # BOTONES UNIFICADOS
                with col1:
                    st.download_button(
                        label="Descargar en WORD (Todo en uno)",
                        data=create_word_file(exam, sol),
                        file_name="Examen_Completo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        label="Descargar en PDF (Todo en uno)",
                        data=create_pdf_file(exam, sol),
                        file_name="Examen_Completo.pdf",
                        mime="application/pdf"
                    )
                    
            else: st.warning("Muy poco texto.")

    elif uploaded_file and not api_key:
        st.warning("AIzaSyBIpNbVgVml5rNrfsIMT0uDnUUhM-GUlLQ")

    # --- PIE DE PÁGINA ---
    st.markdown("---")
    st.markdown("<center><small>Desarrollado por Charles J. Martín M. Arredondo - 2026</small></center>", unsafe_allow_html=True)